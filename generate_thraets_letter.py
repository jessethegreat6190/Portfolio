from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Header
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run('Jondu Jesse')
run.bold = True
run.font.size = Pt(14)
p = doc.add_paragraph('Kampala, Uganda')
p = doc.add_paragraph('jesseford6190@gmail.com | +256 754 490 237')
p = doc.add_paragraph('GitHub: github.com/jessethegreat6190')
p = doc.add_paragraph('Portfolio: https://portfolio-5b977.web.app/')
doc.add_paragraph()

# Date
p = doc.add_paragraph('June 12, 2026')
doc.add_paragraph()

# Recipient
p = doc.add_paragraph('Hiring Manager')
p = doc.add_paragraph('Thraets')
p = doc.add_paragraph('Kampala, Uganda')
doc.add_paragraph()

# Subject
p = doc.add_paragraph()
run = p.add_run('RE: Application for Junior Developer/Designer Position')
run.bold = True
doc.add_paragraph()

# Body
paragraphs = [
    'Dear Hiring Manager,',
    
    'I am writing to apply for the Junior Developer/Designer position at Thraets. As a developer-designer hybrid who builds with AI as a creative partner, I am drawn to Thraets\' mission of combating digital threats against truth, democracy, and social unity. I believe my blend of technical skills, design sensibilities, and passion for civic tech can contribute meaningfully to your work.',
    
    'I am a vibe coder — I use AI tools (Claude, ChatGPT, GitHub Copilot) as force multipliers to rapidly prototype, build, and ship functional digital systems. This approach allows me to move from concept to working product in days, not weeks, while maintaining quality and attention to detail. I have applied this workflow to build:',
]

projects = [
    'A biometric enrollment and identity verification platform (UPDMS) with fingerprint capture, webcam photo acquisition, and OCR document scanning — built end-to-end using AI-assisted development',
    'A full-featured non-profit website with event registration, donation integration, and WhatsApp connectivity',
    'A digital check-in system with QR codes, automated email confirmations, and real-time attendee dashboards',
    'A personal portfolio showcasing 100+ graphic design works with dark mode, animations, and category filtering',
]

closing = [
    'I bring strong frontend development skills (HTML, CSS, JavaScript, React), design capabilities (graphic design, UI/UX, brand identity), and the ability to work across the full stack with PHP, Python, and Node.js. I am also experienced in digital content creation and social media graphics — having delivered over 100 designs for various clients and organizations.',
    
    'What excites me most about Thraets is the opportunity to build tools that empower journalists, human rights defenders, and civil society organizations. I am eager to apply my development and design skills to projects that strengthen democratic processes and combat disinformation in Uganda and across Africa.',
    
    'I am available to start immediately and would welcome the opportunity to discuss how my skills can support Thraets\' mission.',
    
    'Yours sincerely,',
    '',
    'Jondu Jesse',
]

for text in paragraphs:
    p = doc.add_paragraph(text)

for proj in projects:
    doc.add_paragraph(f'• {proj}')

for text in closing:
    p = doc.add_paragraph(text)

doc.save('C:\\xampp\\htdocs\\Portfolio\\application-letter-thraets.docx')
print('Letter saved: application-letter-thraets.docx')
