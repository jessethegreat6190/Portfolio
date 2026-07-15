from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Styles
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)

# Header
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('JONDU JESSE')
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0, 51, 102)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Kampala, Uganda | jesseford6190@gmail.com | +256 754 490 237')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(100, 100, 100)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('GitHub: github.com/jessethegreat6190 | Portfolio: portfolio-5b977.web.app')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph().add_run('').font.size = Pt(6)

# PROFESSIONAL SUMMARY
p = doc.add_paragraph()
run = p.add_run('PROFESSIONAL SUMMARY')
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph('Creative and resourceful Junior Developer & Designer who builds functional, user-focused digital solutions using AI-assisted development workflows. Combines web development skills (HTML, CSS, JavaScript, React, PHP, Python) with graphic design and UI/UX sensibilities to rapidly prototype and ship working products. Passionate about civic tech, digital rights, and using technology to strengthen democratic processes. Experienced in building complete systems — from biometric enrollment platforms to registration portals — using modern AI tools (Claude, ChatGPT, Copilot) as force multipliers.')

# SKILLS
p = doc.add_paragraph()
run = p.add_run('SKILLS & TECHNOLOGIES')
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0, 51, 102)

skills = [
    ('AI-Assisted Development', 'Vibe coding with Claude, ChatGPT, GitHub Copilot — rapid prototyping, debugging, full-stack generation. Build production systems in hours, not weeks.'),
    ('Frontend Development', 'HTML5, CSS3, JavaScript, React.js, Tailwind CSS, responsive design, accessibility basics'),
    ('Backend & Databases', 'PHP, MySQL, Node.js, Firebase, REST API integration'),
    ('Design & UI/UX', 'Graphic design (Canva, Photoshop), UI mockups, typography, brand identity, Figma basics'),
    ('Digital Content', 'Social media graphics, banners, posters, video thumbnails, 100+ published designs'),
    ('Tools & Workflow', 'Git/GitHub, VS Code, Vite, Flutter/Dart (basic), Python (Pandas, NumPy), Linux CLI'),
]

for skill_name, desc in skills:
    p = doc.add_paragraph()
    run = p.add_run(f'{skill_name}: ')
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(desc)
    run.font.size = Pt(11)

# PROJECTS
p = doc.add_paragraph()
run = p.add_run('KEY PROJECTS')
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0, 51, 102)

projects = [
    ('UPDMS — Biometric Enrollment System', 'Built a complete identity verification platform with fingerprint capture (browser API), webcam photo acquisition, and OCR ID scanning (Tesseract.js). Features real-time data validation, PDF report generation, and secure data handling. Built end-to-end using AI-assisted development.'),
    ('Bulk SMS Platform', 'Developed a React + Vite web application for sending bulk SMS messages. Features include a dashboard, CSV contact upload, delivery reports, modem status monitoring, queue management, and settings panel. Built with Tailwind CSS.'),
    ('Dairy Management System', 'Created a full-stack PHP/MySQL dairy farm management system with role-based access (worker, manager, admin). Manages animals, milk production, feeding schedules, health checks, expenses, sales, and reporting.'),
    ('Registration & Check-in System', 'Created a digital registration and check-in platform with QR code generation, automated email confirmations, waitlist management, and real-time attendee dashboard.'),
    ('Portfolio Platform', 'Built a personal portfolio showcasing 100+ graphic design works, projects, and client testimonials. Features dark mode, animations, search, and category filtering.'),
    ('Admin Login System', 'Developed a PHP/MySQL user authentication system with secure registration, login/logout, and role-based access control. Features separate admin and user dashboards, password hashing, and session management.'),
    ('Access Internet — Payment Portal', 'Built a PHP-based pay-as-you-go internet package purchasing system with Airtel Money and MTN Mobile Money integration. Displays time-based and data-based packages with real-time payment processing.'),
    ('College Event Management', 'Designed a dual-panel HTML/CSS system with an admin dashboard for managing events and a student portal for viewing and registering. Features event listings, scheduling, and role-based views.'),
    ('Data Collection Platform', 'Created a Firebase-powered job application data-collection form that submits applicant details (name, email, phone, position, experience) to Firestore. Features real-time data persistence and a clean responsive UI.'),
]

for title, desc in projects:
    p = doc.add_paragraph()
    run = p.add_run(f'{title} — ')
    run.bold = True
    font.size = Pt(11)
    run = p.add_run(desc)
    run.font.size = Pt(11)

# WORK EXPERIENCE
p = doc.add_paragraph()
run = p.add_run('WORK EXPERIENCE')
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0, 51, 102)

p = doc.add_paragraph()
run = p.add_run('Freelance Developer & Designer')
run.bold = True
run.font.size = Pt(11)
p = doc.add_paragraph('Jan 2024 — Present | Kampala, Uganda')
p.paragraph_format.space_after = Pt(2)
items = [
    'Develop web applications, landing pages, and digital systems for small businesses and non-profits',
    'Use AI-assisted workflows to accelerate development — reducing delivery timelines by 60%+',
    'Design brand identities, social media graphics, marketing collateral (100+ designs delivered)',
    'Manage end-to-end project delivery: requirements gathering, prototyping, development, deployment',
]
for item in items:
    doc.add_paragraph(f'• {item}')

p = doc.add_paragraph()
run = p.add_run('IT Support & Systems Developer')
run.bold = True
run.font.size = Pt(11)
p = doc.add_paragraph('Jan 2024 — Present | Kampala, Uganda')
p.paragraph_format.space_after = Pt(2)
items = [
    'Built biometric enrollment systems for identity verification projects',
    'Provided technical support and equipment maintenance in customer-facing environments',
    'Managed sensitive data with strict confidentiality protocols',
]
for item in items:
    doc.add_paragraph(f'• {item}')

# EDUCATION
p = doc.add_paragraph()
run = p.add_run('EDUCATION')
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0, 51, 102)

p = doc.add_paragraph()
run = p.add_run('Diploma in Information Technology')
run.bold = True
run.font.size = Pt(11)
doc.add_paragraph('ISBAT University, Kampala | Jan 2023 — Jun 2026 (Final semester completed)')

# ADDITIONAL
p = doc.add_paragraph()
run = p.add_run('WHAT SETS ME APART')
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0, 51, 102)

items = [
    'Vibe coder: I use AI as a creative partner to build, debug, and ship faster — turning ideas into working products rapidly',
    'Full-strain thinker: I understand both code and design — bridging the gap between how something works and how it looks',
    'Civic-tech minded: passionate about using technology to strengthen democracy, combat disinformation, and support civil society',
    'Self-taught & curious: continuously learning new tools and methodologies to stay effective',
]
for item in items:
    doc.add_paragraph(f'• {item}')

# REFERENCES
p = doc.add_paragraph()
run = p.add_run('REFERENCES')
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0, 51, 102)

p = doc.add_paragraph()
run = p.add_run('Professional Referees')
run.bold = True
run.font.size = Pt(11)
doc.add_paragraph('1. Kato Kenneth – Lecturer, ISBAT University | Tel: 0760 220 536')
doc.add_paragraph('2. Chitra G. Sasi – Lecturer, ISBAT University | Tel: 0754 411 927')

p = doc.add_paragraph()
run = p.add_run('Character Referees')
run.bold = True
run.font.size = Pt(11)
doc.add_paragraph('1. Nanyanzi Pamela – Advocate & Church Chairman | Tel: 0779 827 270')
doc.add_paragraph('2. Nangobi Hellen Hilda – Church Member | Tel: 0754 007 779')

doc.save('C:\\xampp\\htdocs\\Portfolio\\Jondu_Jesse_CV_Thraets.docx')
print('CV saved: Jondu_Jesse_CV_Thraets.docx')
